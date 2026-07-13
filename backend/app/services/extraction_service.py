"""
Syllabus extraction service — v2 with layout-aware PDF/DOCX extraction.

Primary path  : configured light model (settings.generation_light_model)
                → structured JSON → validated by Pydantic AIExtraction
Fallback path : regex heuristics (no API key / quota exceeded / timeout)

Enhancements over v1:
- Layout-aware PyMuPDF extraction (font-size clustering for H1/H2 detection)
- Scanned PDF detection → raises ScannedPDFError (HTTP 422)
- Encrypted PDF detection → raises EncryptedPDFError (HTTP 422)
- DOCX heading-style extraction
- Section keyword classifier to route text blocks to the right field
- Three regex patterns per field with per-result confidence tagging
- GPT prompt includes detected document structure when available
"""

import io
import re
import json
import logging
from collections import Counter
from app.schemas.syllabus import (
    SyllabusExtraction, AIExtraction, ExtractedField,
    ExtractedCO, ExtractedUnit, ExtractedTopic, ExtractedSubtopic,
)

logger = logging.getLogger(__name__)

VALID_BLOOM = {"Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create"}


# ── Custom exceptions ──────────────────────────────────────────────────────────

class ScannedPDFError(ValueError):
    """Raised when the PDF appears to be a scanned image with no selectable text."""


class EncryptedPDFError(ValueError):
    """Raised when the PDF is password-protected."""


def _light_model() -> str:
    """The light generation model (structured extraction)."""
    try:
        from app.core.config import settings
        return settings.generation_light_model
    except Exception:  # pragma: no cover
        return "gpt-5.4-nano"  # mirror the config default


def _ocr_model() -> str:
    """OCR needs a model with confirmed image input. OCR_MODEL pins it
    explicitly; otherwise the light model is assumed multimodal."""
    try:
        from app.core.config import settings
        return settings.ocr_model or settings.generation_light_model
    except Exception:  # pragma: no cover
        return "gpt-5.6-luna"  # mirror the config default


# ── Scanned / encrypted detection ─────────────────────────────────────────────

def _is_scanned(doc) -> bool:
    """Return True if the first 3 pages collectively have fewer than 50 chars/page."""
    sample_pages = list(doc)[:min(3, len(doc))]
    chars = sum(len(page.get_text()) for page in sample_pages)
    return chars < 50 * len(sample_pages)


def _is_encrypted(content: bytes) -> bool:
    try:
        import fitz
        doc = fitz.open(stream=content, filetype="pdf")
        return doc.needs_pass
    except Exception:
        return False


# ── Vision OCR for scanned PDFs ────────────────────────────────────────────────
#
# A large share of real Indian university syllabi are scans. Instead of
# rejecting them, transcribe each page with a vision model and feed the text
# through the normal extraction paths (5-prompt pipeline / single-prompt / regex)
# unchanged. Returns "" when no key is configured or transcription fails, so
# callers fall back to ScannedPDFError exactly as before.

_OCR_MAX_PAGES = 10   # syllabi are short; caps cost on accidental big uploads
_OCR_DPI = 150

_OCR_INSTRUCTION = (
    "Transcribe ALL text on this scanned syllabus page exactly as printed. "
    "Preserve line breaks, headings, numbering, and list structure; render "
    "tables row by row with cells separated by ' | '. Output ONLY the "
    "transcription — no commentary, no markdown fences."
)


def ocr_scanned_pdf(content: bytes, max_pages: int = _OCR_MAX_PAGES) -> str:
    """Page-by-page vision transcription of a scanned PDF. Best-effort: any
    failure returns "" and the caller keeps its existing scanned-PDF handling."""
    try:
        import base64
        import fitz
        from openai import OpenAI
        from app.core.config import settings

        if not settings.openai_api_key:
            return ""
        client = OpenAI(api_key=settings.openai_api_key)
        doc = fitz.open(stream=content, filetype="pdf")
        pages: list[str] = []
        from app.services import llm_compat
        for page in list(doc)[:max_pages]:
            pix = page.get_pixmap(dpi=_OCR_DPI)
            b64 = base64.b64encode(pix.tobytes("png")).decode()
            resp = llm_compat.create_chat_completion(
                client, model=_ocr_model(),
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": _OCR_INSTRUCTION},
                    {"type": "image_url",
                     "image_url": {"url": f"data:image/png;base64,{b64}"}},
                ]}],
                temperature=0,
                max_tokens=4096,
            )
            pages.append(resp.choices[0].message.content or "")
        text = "\n\n".join(p.strip() for p in pages if p.strip())
        logger.info("vision OCR transcribed %d page(s), %d chars", len(pages), len(text))
        return text
    except Exception:
        logger.warning("vision OCR failed", exc_info=True)
        return ""


def extract_text_pdf(content: bytes) -> str:
    """Raw text for the syllabus pipeline: layout text for digital PDFs, vision
    OCR for scanned ones.

    Raises EncryptedPDFError / ScannedPDFError (only when OCR is unavailable
    or produced nothing) / ValueError like the structured extractor."""
    try:
        import fitz
    except ImportError:
        raise ValueError("PyMuPDF (fitz) is not installed. Run: pip install pymupdf")
    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ValueError(f"Could not open PDF: {exc}") from exc
    if doc.needs_pass:
        raise EncryptedPDFError("PDF is password-protected. Remove the password and re-upload.")

    if not _is_scanned(doc):
        flat_text, _ = _pdf_to_layout(content)
        if flat_text.strip():
            return flat_text
        flat_text = "\n".join(page.get_text() for page in doc)
        if flat_text.strip():
            return flat_text

    ocr_text = ocr_scanned_pdf(content)
    if ocr_text.strip():
        return ocr_text
    raise ScannedPDFError(
        "This appears to be a scanned PDF and OCR is unavailable. "
        "Please use a digital PDF or paste the text manually."
    )


# ── Runtime module-tree coverage validator ────────────────────────────────────
#
# A correct extraction prompt can still occasionally drop content on a given
# run (observed: File-System Internals items vanished from one otherwise-perfect
# extraction). The golden-set eval catches this per prompt change; this
# validator catches it per UPLOAD: segment the raw text by unit, token-compare
# each unit's body against its extracted tree, and flag under-covered units for
# the faculty review screen. Deterministic, offline, advisory.

_COVERAGE_THRESHOLD = 0.70

_COV_STOP = {
    "with", "from", "that", "this", "their", "using", "various", "different",
    "between", "other", "into", "based", "unit", "units", "module", "modules",
    "hours", "hour", "chapter", "topics", "subtopics",
}

_UNIT_HEAD_PAT = re.compile(r'^(?:UNIT|MODULE)\s*[-–—]?\s*([IVX\d]+)\b.*$',
                            re.IGNORECASE | re.MULTILINE)
# The heading may be alone on its line ("TEXT BOOKS") or carry the list inline
# ("Text Books: 1) Miller and Freund's ..."). Inline content requires the colon
# so prose that merely starts with "References ..." doesn't end the unit.
_SECTION_END_PAT = re.compile(
    r'^\s*(?:text\s*books?|reference\s*books?|references|bibliography|'
    r'online\s+(?:learning\s+)?resources?|e-?resources?|web\s+resources?)'
    r'\s*(?::.*)?$',
    re.IGNORECASE | re.MULTILINE)

_ROMAN_TO_INT = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
                 'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10}


def _unit_segments(text: str) -> list[tuple[int, str]]:
    """(unit_number, body_text) for each unit heading found in the raw text.
    A unit's body runs to the next unit heading or the books/resources section."""
    matches = list(_UNIT_HEAD_PAT.finditer(text))
    segments: list[tuple[int, str]] = []
    for i, m in enumerate(matches):
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end]
        em = _SECTION_END_PAT.search(body)
        if em:
            body = body[:em.start()]
        raw = m.group(1).upper()
        num = _ROMAN_TO_INT.get(raw) or (int(raw) if raw.isdigit() else i + 1)
        segments.append((num, body))
    return segments


def _coverage_tokens(body: str, full_text: str) -> set[str]:
    """Distinctive content tokens of a unit body. Lines that repeat verbatim
    elsewhere in the document (page headers/footers) and digit-heavy lines
    (L/T/P/C rows, page numbers) are dropped first."""
    line_counts = Counter(l.strip() for l in full_text.splitlines() if l.strip())
    tokens: set[str] = set()
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped or line_counts[stripped] > 1:
            continue
        words = stripped.split()
        digitish = sum(1 for w in words if len(w) <= 2 or any(c.isdigit() for c in w))
        if words and digitish / len(words) > 0.5:
            continue
        for w in re.findall(r"[a-z][a-z\-']{3,}",
                            stripped.casefold().replace("’", "'")):
            w = w.strip("-'").removesuffix("'s")
            if len(w) >= 4 and w not in _COV_STOP:
                tokens.add(w)
    return tokens


def _tree_tokens(unit: dict) -> set[str]:
    """All tokens present in a unit's extracted topics + subtopics."""
    parts: list[str] = []
    for t in unit.get("topics") or []:
        parts.append(str(t.get("title", "")))
        for s in t.get("subtopics") or []:
            parts.append(str(s.get("title", "")) if isinstance(s, dict) else str(s))
    blob = " ".join(parts).casefold().replace("’", "'")
    out = set()
    for source in (blob, blob.replace("-", " ")):
        for w in re.findall(r"[a-z][a-z\-']{3,}", source):
            out.add(w)
            out.add(w.strip("-'").removesuffix("'s"))
    return out


def module_tree_coverage(text: str, units: list[dict],
                         threshold: float = _COVERAGE_THRESHOLD) -> dict:
    """Per-unit coverage report: which share of each unit's distinctive source
    tokens made it into the extracted tree, with the missing terms named.
    `units` is ai_extraction-shaped ({unit_number, topics:[{title, subtopics}]});
    subtopics may be strings or {title} dicts."""
    by_number = {int(u.get("unit_number") or 0): u for u in units}
    report_units = []
    flagged: list[int] = []
    total_found = total_tokens = 0
    for num, body in _unit_segments(text):
        unit = by_number.get(num)
        want = _coverage_tokens(body, text)
        have = _tree_tokens(unit) if unit else set()
        found = want & have
        missing = sorted(want - have)
        cov = len(found) / len(want) if want else 1.0
        total_found += len(found)
        total_tokens += len(want)
        is_flagged = cov < threshold or unit is None
        if is_flagged:
            flagged.append(num)
        report_units.append({
            "unit_number": num,
            "coverage": round(cov, 3),
            "extracted": unit is not None,
            "missing_terms": missing[:15],
            "flagged": is_flagged,
        })
    return {
        "overall_coverage": round(total_found / total_tokens, 3) if total_tokens else 1.0,
        "threshold": threshold,
        "units": report_units,
        "flagged_units": flagged,
    }


# ── Section keyword classifier ─────────────────────────────────────────────────

SECTION_KEYWORDS: dict[str, list[str]] = {
    "cos": [
        "course outcome", "learning outcome", "co ", "students will be able",
        "after completing", "course objective", "upon completion",
    ],
    "hours": [
        "l:", "t:", "p:", "lecture", "tutorial", "credits", "l-t-p", "ltp",
        "total hours", "contact hours",
    ],
    "books": [
        "reference", "textbook", "bibliography", "recommended reading",
    ],
    "assessment": [
        "assessment", "evaluation", "marks", "ese", "cia", "internal",
        "external", "end semester",
    ],
    "lab": [
        "laboratory", "practical", "lab", "experiment", "hands-on",
    ],
}


def _classify_section(text: str) -> str | None:
    lower = text.lower()
    for section, keywords in SECTION_KEYWORDS.items():
        if any(kw in lower for kw in keywords):
            return section
    return None


# ── PDF layout extraction ──────────────────────────────────────────────────────

def _pdf_to_layout(content: bytes) -> tuple[str, list[dict]]:
    """
    Returns (flat_text, layout_units).

    layout_units is a list of dicts:
        { "title": str, "topics": [str], "raw_text": str }

    Uses font-size clustering to identify H1 (unit titles) and H2 (topic titles).
    Falls back to returning a single empty unit list when no clear hierarchy exists.
    """
    try:
        import fitz
    except ImportError:
        logger.error("PyMuPDF (fitz) not installed")
        return "", []

    doc = fitz.open(stream=content, filetype="pdf")

    # Collect all text spans with their font sizes
    spans: list[dict] = []
    for page in doc:
        blocks = page.get_text("dict")["blocks"]  # type: ignore[attr-defined]
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span["text"].strip()
                    if text:
                        spans.append({
                            "text": text,
                            "size": round(span["size"], 1),
                            "bold": bool(span["flags"] & 2**4),
                        })

    if not spans:
        return "", []

    # Determine body size (most frequent font size)
    size_counts = Counter(s["size"] for s in spans)
    body_size: float = size_counts.most_common(1)[0][0]

    all_sizes = sorted(set(s["size"] for s in spans), reverse=True)

    # H1: largest size that is at least 15% bigger than body
    h1_size: float | None = next(
        (s for s in all_sizes if s > body_size * 1.15), None
    )

    # H2: sizes between body and H1 (exclusive)
    h2_sizes: set[float] = set(
        s for s in all_sizes
        if body_size < s < (h1_size if h1_size else body_size * 2)
    )

    # Build structure
    flat_lines: list[str] = []
    units: list[dict] = []
    current_unit: dict | None = None
    current_topic_texts: list[str] = []

    for span in spans:
        size = span["size"]
        text = span["text"]
        flat_lines.append(text)

        if h1_size is not None and abs(size - h1_size) < 0.5:
            # Flush previous unit
            if current_unit is not None:
                current_unit["raw_text"] = " ".join(current_topic_texts)
                units.append(current_unit)
            current_unit = {"title": text, "topics": [], "raw_text": ""}
            current_topic_texts = []

        elif h2_sizes and size in h2_sizes:
            if current_unit is not None:
                current_unit["topics"].append(text)
            current_topic_texts.append(text)

        else:
            current_topic_texts.append(text)

    # Flush the last unit
    if current_unit is not None:
        current_unit["raw_text"] = " ".join(current_topic_texts)
        units.append(current_unit)

    flat_text = "\n".join(flat_lines)
    return flat_text, units


# ── DOCX layout extraction ─────────────────────────────────────────────────────

def _docx_to_layout(content: bytes) -> tuple[str, list[dict]]:
    """
    Returns (flat_text, layout_units) from a DOCX file.
    Uses paragraph heading styles (Heading 1/2/3) for hierarchy detection.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed")
        return "", []

    doc = Document(io.BytesIO(content))
    flat_lines: list[str] = []
    units: list[dict] = []
    current_unit: dict | None = None
    current_topic_texts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        flat_lines.append(text)

        style = para.style.name if para.style else ""

        if "Heading 1" in style:
            if current_unit is not None:
                current_unit["raw_text"] = " ".join(current_topic_texts)
                units.append(current_unit)
            current_unit = {"title": text, "topics": [], "raw_text": ""}
            current_topic_texts = []

        elif "Heading 2" in style:
            if current_unit is not None:
                current_unit["topics"].append(text)
            current_topic_texts.append(text)

        elif "Heading 3" in style:
            if current_unit is not None:
                current_unit["topics"].append(text)
            current_topic_texts.append(text)

        else:
            current_topic_texts.append(text)

    if current_unit is not None:
        current_unit["raw_text"] = " ".join(current_topic_texts)
        units.append(current_unit)

    flat_text = "\n".join(flat_lines)
    return flat_text, units


# ── Regex field extractors (3 patterns each, with confidence) ──────────────────

def _extract_hours(text: str) -> tuple[str, str]:
    """
    Try to extract total contact hours from text.
    Returns (value_str, confidence).
    """
    # Pattern 1: explicit L/T/P table — "L: 3  T: 1  P: 0" or "L-3 T-1 P-0"
    m = re.search(r'L\s*[:\-]\s*(\d+)\s*T\s*[:\-]\s*(\d+)\s*P\s*[:\-]\s*(\d+)', text, re.I)
    if m:
        total = int(m.group(1)) + int(m.group(2)) + int(m.group(3))
        return str(total), "High"

    # Pattern 1b: numeric triple "3-1-0" or "3–1–0"
    m = re.search(r'(\d+)\s*[-–]\s*(\d+)\s*[-–]\s*(\d+)', text)
    if m:
        total = int(m.group(1)) + int(m.group(2)) + int(m.group(3))
        return str(total), "High"

    # Pattern 2: explicit total hours label
    m = re.search(r'total\s*(?:contact\s*)?hours?\s*[:\-]\s*(\d+)', text, re.I)
    if m:
        return m.group(1), "High"

    # Pattern 3: credit count — multiply by 15 (heuristic)
    m = re.search(r'(\d+)\s*credits?', text, re.I)
    if m:
        return str(int(m.group(1)) * 15), "Low"

    # Pattern 4: hours per week
    m = re.search(r'(\d+)\s*hours?\s*(?:per\s*week)', text, re.I)
    if m:
        return str(int(m.group(1)) * 15), "Low"

    return "", "Missing"


def _extract_cos(text: str) -> tuple[list[str], str]:
    """
    Try to extract CO statements.
    Returns ([co_text, ...], confidence).
    """
    # Pattern 1: explicit "CO1:", "CO 2 -", etc.
    cos = re.findall(r'CO\s*\d+\s*[:\-]?\s*(.+)', text, re.I)
    clean = [c.strip() for c in cos if len(c.strip()) > 10]
    if clean:
        return clean, "High"

    # Pattern 2: numbered list inside a CO section
    cos = re.findall(r'(?:^|\n)\s*\d+[\.\)]\s*(.{20,})', text)
    clean = [c.strip() for c in cos]
    if clean:
        return clean, "Low"

    # Pattern 3: bullet points
    cos = re.findall(r'(?:^|\n)\s*[•●▪\-]\s*(.{20,})', text)
    clean = [c.strip() for c in cos]
    if clean:
        return clean, "Low"

    return [], "Missing"


def _extract_assessment(text: str) -> tuple[dict | None, str]:
    """
    Try to extract assessment scheme (ESE %, CIA %).
    Returns (dict_or_None, confidence).
    """
    # Pattern 1: explicit ESE / CIA labels
    ese = re.search(r'(?:ESE|End\s*Semester\s*Exam(?:ination)?)[^\d]*(\d+)\s*%?', text, re.I)
    cia = re.search(r'(?:CIA|Internal|CIE|Continuous\s*Internal)[^\d]*(\d+)\s*%?', text, re.I)
    if ese or cia:
        return {
            "ese_pct": int(ese.group(1)) if ese else None,
            "cia_pct": int(cia.group(1)) if cia else None,
        }, "High"

    # Pattern 2: percentage after generic marks keywords
    marks = re.findall(r'([\w][\w\s]+?)\s*[:\-]\s*(\d+)\s*(?:marks|%)', text, re.I)
    if marks:
        result = {k.strip().lower(): int(v) for k, v in marks[:4]}
        return result, "Low"

    # Pattern 3: fraction out-of pattern  e.g. "80/100"
    frac = re.findall(r'(\d+)\s*/\s*100', text)
    if len(frac) >= 2:
        return {"component_1": int(frac[0]), "component_2": int(frac[1])}, "Low"

    return None, "Missing"


# ── AI extraction (primary) ───────────────────────────────────────────────────

_GPT_SYSTEM = """You are a university syllabus parser specialising in Indian engineering university syllabi (JNTU, Anna University, VTU, KTU, etc.).

Extract the full structure and return ONLY valid JSON — no markdown, no explanation.

CRITICAL — Indian syllabus format rules:
Indian syllabi use a very specific structure. You MUST follow these rules exactly:

1. UNITS are announced by lines like "UNIT – I (10 Hours)", "UNIT I", "MODULE 1 (9 Hrs)" etc.
   If the syllabus gives the unit an explicit name on or next to that line
   (e.g. "UNIT – I: Operating System Fundamentals (10 Hours)"), use that name as the
   unit title. If no explicit name exists, set the title to "Unit I" / "Unit II" etc.
   NEVER promote a topic to the unit title — the unit title must never be identical
   to any topic title inside that unit.

2. TOPICS are bold phrases followed by a colon, e.g. "Operating Systems Overview:" or "Processes:"
   Each bold-phrase-before-colon is ONE topic.

3. SUBTOPICS are the comma-separated items listed AFTER the colon on the same line/paragraph.
   e.g. "Processes: Process Concept, Process scheduling, Operations on processes, Inter-process communication"
   → topic title = "Processes"
   → subtopics = ["Process Concept", "Process scheduling", "Operations on processes", "Inter-process communication"]
   Keep parenthetical enumerations attached to their item — "Scheduling algorithms
   (FCFS, SJF, Round Robin)" is ONE subtopic, not three.

4. DO NOT treat each comma-separated item as a separate topic. They are subtopics.
   A phrase-list line with no heading of its own continues the PREVIOUS topic —
   append its items to that topic's subtopics.

5. A unit typically has 2–4 topics. If you find more than 6 topics in one unit, you are almost certainly
   splitting subtopics into topics — merge them back under their bold heading.
   Group sibling sections of one broad concept ("File-System Interface",
   "File-System Implementation", "File-System Internals") into ONE topic.

Return exactly this shape:
{
  "course_name": "string",
  "course_code": "string",
  "credits": integer,
  "semester": "string",
  "regulation": "string",
  "total_hours": integer,
  "lab_flag": boolean,
  "course_outcomes": [
    { "co_number": 1, "text": "full CO statement", "bloom_level": "one of Remember|Understand|Apply|Analyze|Evaluate|Create" }
  ],
  "units": [
    {
      "unit_number": 1,
      "title": "explicit unit name if the syllabus gives one, else 'Unit I' — never a topic title",
      "hours": integer,
      "topics": [
        {
          "title": "topic name — the bold heading before the colon",
          "bloom_level": "inherit from the CO this topic supports — do NOT default all to Understand",
          "subtopics": [
            { "title": "each comma-separated item after the colon" }
          ]
        }
      ]
    }
  ],
  "reference_books": ["Author, Title, Edition"]
}

Additional rules:
- For topic bloom_level: first check if the syllabus explicitly states a bloom verb for that topic.
  If not (most Indian syllabi list topics as bare nouns), judge the level the topic is genuinely
  TAUGHT at from its own subtopics, using the CO it supports as the CEILING — never the default:
  survey/definitional topics → Remember/Understand; mechanism/single-technique topics →
  Understand/Apply; problem-solving/implementation topics → Apply; comparison/analysis topics
  → Analyze; design/architecture topics → Evaluate/Create. An introductory topic under an
  Analyze-level CO is still Understand. Do NOT default everything to Understand, and do NOT
  copy the CO's level onto every topic it covers.
- CO bloom_level inference from action verbs: list/recall→Remember, explain/describe→Understand,
  implement/use/apply→Apply, compare/examine→Analyze, evaluate/judge→Evaluate, design/create→Create.
- If a field is not present, use empty string / 0 / [] as appropriate.
- lab_flag: true if the syllabus mentions labs, practicals, or P credits > 0.
- course_outcomes: extract verbatim CO statements if present; otherwise derive one per unit.
- Return ONLY the JSON object. No extra text, no markdown fences."""


def _parse_indian_syllabus_structure(text: str) -> list[dict]:
    """
    Pre-process Indian university syllabus text to extract the unit→topic→subtopic
    hierarchy before sending to GPT.

    Indian format:
      UNIT – I (10 Hours)
      Operating Systems Overview: Introduction, OS functions, OS operations, ...
      System Structures: OS Services, User interface, system calls, ...

    Returns a list of units:
      [{"unit_number": 1, "title": "...", "hours": 10,
        "topics": [{"title": "...", "subtopics": ["...", ...]}]}]
    """
    units: list[dict] = []
    current_unit: dict | None = None

    # Unit heading pattern: UNIT – I (10 Hours) / UNIT I / MODULE 1 (9 Hrs)
    unit_pat = re.compile(
        r'^(?:UNIT|MODULE)\s*[-–]?\s*([IVX\d]+)\s*(?:\(\s*(\d+)\s*[Hh]ours?\s*\))?',
        re.IGNORECASE
    )
    # Topic line: bold heading followed by colon then subtopics.
    # The topic title must not contain a comma (to avoid matching mid-sentence colons).
    topic_pat = re.compile(r'^([^:,\n]{3,70}):\s*(.{5,})$')

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        # Check for unit heading
        um = unit_pat.match(line)
        if um:
            if current_unit:
                units.append(current_unit)
            # Convert Roman to int for unit_number
            roman = um.group(1).upper()
            roman_map = {'I':1,'II':2,'III':3,'IV':4,'V':5,'VI':6,
                         'VII':7,'VIII':8,'IX':9,'X':10}
            unit_num = roman_map.get(roman) or (int(roman) if roman.isdigit() else len(units) + 1)
            hours = int(um.group(2)) if um.group(2) else 0
            if not hours:
                # "UNIT – I: Name (10 Hours)" — hours after the explicit name.
                hm = re.search(r'\(\s*(\d+)\s*[Hh](?:ou)?rs?\.?\s*\)', line)
                hours = int(hm.group(1)) if hm else 0
            # Explicit unit name on the heading line itself, e.g.
            # "UNIT – I: Operating System Fundamentals (10 Hours)".
            rest = re.sub(r'\(\s*\d+\s*[Hh](?:ou)?rs?\.?\s*\)', '', line[um.end():])
            rest = rest.strip(' \t-–:—')
            title = rest if 3 <= len(rest) <= 80 else ""
            current_unit = {"unit_number": unit_num, "title": title, "hours": hours, "topics": []}
            continue

        # Check for topic line (bold-heading: subtopics)
        tm = topic_pat.match(line)
        if tm and current_unit is not None:
            topic_title = tm.group(1).strip().rstrip(':')
            subtopics_raw = tm.group(2).strip()
            # Split subtopics by comma, filter noise
            subtopics = [s.strip() for s in subtopics_raw.split(',') if len(s.strip()) > 2]
            current_unit["topics"].append({"title": topic_title, "subtopics": subtopics})
        elif current_unit is not None and not current_unit["title"] and len(line) < 80:
            # Short standalone line right after the unit heading = the unit's
            # explicit name printed on its own line.
            current_unit["title"] = line

    if current_unit:
        units.append(current_unit)

    # A unit with no explicit name gets "Unit N" — never a promoted topic title
    # (P1 pipeline rule; both extractors must agree so the same document cannot
    # produce two different unit trees depending on which path ran).
    _ROMAN = {1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V',
              6: 'VI', 7: 'VII', 8: 'VIII', 9: 'IX', 10: 'X'}
    for u in units:
        if not u["title"]:
            u["title"] = f"Unit {_ROMAN.get(u['unit_number'], u['unit_number'])}"

    return units


def _ai_extract(text: str, layout_units: list[dict] | None = None) -> "AIExtraction | None":
    """
    Call the configured light model (settings.generation_light_model) and
    validate the response with Pydantic.
    Returns None on any failure (missing key, quota, timeout).

    When layout_units is provided, the document structure is prepended to the
    user message so GPT can use it as a structural guide during extraction.
    """
    try:
        from openai import OpenAI
        from app.core.config import settings

        if not settings.openai_api_key:
            logger.info("No OpenAI key — skipping AI extraction")
            return None

        client = OpenAI(api_key=settings.openai_api_key)

        # ~10k tokens of input — the old 12k-char cap (~3k tokens) silently
        # dropped the later units of long syllabi.
        truncated = text[:40_000]

        # Try Indian syllabus pre-parser first (most accurate for this format)
        indian_units = _parse_indian_syllabus_structure(truncated)

        if indian_units and any(u["topics"] for u in indian_units):
            structure_hint = (
                "PARSED DOCUMENT STRUCTURE (authoritative — use exactly this hierarchy):\n"
                "Each item below shows: Unit → Topic → [subtopics].\n"
                "DO NOT split subtopics into separate topics.\n\n"
            )
            for u in indian_units:
                structure_hint += f"UNIT {u['unit_number']} ({u['hours']} hrs): {u['title']}\n"
                for t in u["topics"]:
                    subs = ", ".join(t["subtopics"][:8])
                    structure_hint += f"  TOPIC: {t['title']}\n"
                    if subs:
                        structure_hint += f"    Subtopics: {subs}\n"
            user_msg = structure_hint + "\n\nFull syllabus text (for course metadata, COs, books):\n\n" + truncated
        elif layout_units:
            structure_hint = "Detected document structure (use this to guide extraction):\n"
            for i, u in enumerate(layout_units):
                structure_hint += f"Unit {i + 1}: {u['title']}\n"
                for t in u.get("topics", []):
                    structure_hint += f"  - {t}\n"
            user_msg = structure_hint + "\n\nFull syllabus text:\n\n" + truncated
        else:
            user_msg = f"Syllabus text:\n\n{truncated}"

        from app.services import llm_compat
        response = llm_compat.create_chat_completion(
            client, model=_light_model(),
            messages=[
                {"role": "system", "content": _GPT_SYSTEM},
                {"role": "user",   "content": user_msg},
            ],
            temperature=0,
            max_tokens=8192,  # a dense multi-unit tree overflows 4096 output tokens
            response_format={"type": "json_object"},
        )

        if getattr(response.choices[0], "finish_reason", None) == "length":
            logger.warning("AI extraction output truncated (finish_reason=length)")

        raw = response.choices[0].message.content or "{}"
        data = json.loads(raw)

        validated = AIExtraction.model_validate(data)
        logger.info(
            "AI extraction OK — %d COs, %d units",
            len(validated.course_outcomes),
            len(validated.units),
        )
        return validated

    except Exception as e:
        logger.warning("AI extraction failed (%s): %s", type(e).__name__, e)
        return None


# ── Regex fallback (secondary) ────────────────────────────────────────────────

def _field(value: object, confidence: str = "High") -> ExtractedField:
    return ExtractedField(value=value, confidence=confidence if value else "Missing")


def _regex_extract(text: str) -> SyllabusExtraction:
    """Full regex-based extraction. Used when AI extraction fails or key is absent."""
    lines = [l.strip() for l in text.splitlines() if l.strip()]

    def find(patterns: list[str]) -> tuple[str, str]:
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                return m.group(1).strip(), "High"
        return "", "Missing"

    course_name, c1 = find([r"course\s*title\s*[:\-]\s*(.+)", r"subject\s*name\s*[:\-]\s*(.+)"])
    course_code, c2 = find([r"course\s*code\s*[:\-]\s*(\w+)", r"subject\s*code\s*[:\-]\s*(\w+)"])
    credits, c3     = find([r"credits?\s*[:\-]\s*(\d+)", r"L\s*[\+\-]\s*T\s*[\+\-]\s*P\s*[:\-]\s*([\d\s\+]+)"])
    semester, c4    = find([r"semester\s*[:\-]\s*(.+)", r"sem\s*[:\-]\s*(.+)"])

    cos, cos_conf = _extract_cos(text)
    if not cos:
        # Additional fallback: "course outcome N:"
        raw = re.findall(r"course\s*outcome\s*\d+\s*[:\-]?\s*(.+)", text, re.IGNORECASE)
        cos = [r.strip() for r in raw if len(r.strip()) > 10]
        cos_conf = "High" if cos else "Missing"

    modules = re.findall(r"(?:module|unit)\s*[:\-\s]?\s*(\d+)\s*[:\-]\s*(.+)", text, re.IGNORECASE)
    modules_list = [f"Module {n}: {t}" for n, t in modules]

    total_hours_str, h_conf = _extract_hours(text)
    if not total_hours_str:
        # legacy single-pattern fallback
        m = re.search(r"(\d+)\s*hours?", text, re.I)
        if m:
            total_hours_str, h_conf = m.group(1), "Low"

    refs = re.findall(r"(?:\d+[\.\)]\s*)([A-Z][^,\n]{10,})", text)
    refs = [r for r in refs if len(r) > 15]
    lab = bool(re.search(r"\blab\b|\bpractical\b|\bP\s*=\s*[1-9]", text, re.IGNORECASE))

    assessment, a_conf = _extract_assessment(text)

    return SyllabusExtraction(
        course_name=_field(course_name, c1),
        course_code=_field(course_code, c2),
        credits=_field(credits, c3),
        semester=_field(semester, c4),
        course_outcomes=_field(cos or None, cos_conf),
        modules=_field(modules_list or None, "High" if modules_list else "Low"),
        total_hours=_field(total_hours_str or None, h_conf),
        reference_books=_field(refs[:10] or None, "High" if refs else "Missing"),
        lab_flag=_field(str(lab), "High"),
        assessment_scheme=_field(assessment, a_conf),
        ai_extraction=None,
    )


# ── AI result → SyllabusExtraction (unified shape) ───────────────────────────

def _ai_to_syllabus(ai: AIExtraction) -> SyllabusExtraction:
    """Convert validated AIExtraction into SyllabusExtraction."""
    return SyllabusExtraction(
        course_name=_field(ai.course_name or None),
        course_code=_field(ai.course_code or None),
        credits=_field(str(ai.credits) if ai.credits else None),
        semester=_field(ai.semester or None),
        course_outcomes=_field(
            [co.text for co in ai.course_outcomes] if ai.course_outcomes else None,
            "High" if ai.course_outcomes else "Missing",
        ),
        modules=_field(
            [f"Unit {u.unit_number}: {u.title}" for u in ai.units] if ai.units else None,
            "High" if ai.units else "Missing",
        ),
        total_hours=_field(str(ai.total_hours) if ai.total_hours else None),
        reference_books=_field(ai.reference_books or None, "High" if ai.reference_books else "Missing"),
        lab_flag=_field(str(ai.lab_flag), "High"),
        assessment_scheme=_field(None, "Missing"),
        ai_extraction=ai,
    )


# ── Re-extract a single field from stored text ────────────────────────────────

def reextract_field(field_name: str, selected_text: str) -> tuple[object, str]:
    """
    Run field-specific re-extraction on selected_text.
    Returns (value, confidence).

    Tries regex first; falls back to GPT for the field if regex misses.
    """
    text = selected_text

    if field_name == "hours":
        value, confidence = _extract_hours(text)
        if not value:
            ai = _ai_extract(text)
            if ai and ai.total_hours:
                return str(ai.total_hours), "Low"
        return value or None, confidence

    elif field_name == "cos":
        value, confidence = _extract_cos(text)
        if not value:
            ai = _ai_extract(text)
            if ai and ai.course_outcomes:
                return [co.text for co in ai.course_outcomes], "Low"
        return value or None, confidence

    elif field_name == "books":
        refs = re.findall(r"(?:\d+[\.\)]\s*)([A-Z][^,\n]{10,})", text)
        refs = [r for r in refs if len(r) > 15]
        if refs:
            return refs[:10], "High"
        # Try broader patterns
        lines = [l.strip() for l in text.splitlines() if l.strip() and len(l.strip()) > 15]
        if lines:
            return lines[:10], "Low"
        return None, "Missing"

    elif field_name == "assessment":
        value, confidence = _extract_assessment(text)
        if not value:
            ai = _ai_extract(text)
            # Assessment isn't directly in AIExtraction schema; return raw regex result
        return value, confidence

    elif field_name == "modules":
        modules = re.findall(r"(?:module|unit)\s*[:\-\s]?\s*(\d+)\s*[:\-]\s*(.+)", text, re.IGNORECASE)
        modules_list = [f"Module {n}: {t.strip()}" for n, t in modules]
        if modules_list:
            return modules_list, "High"
        ai = _ai_extract(text)
        if ai and ai.units:
            return [f"Unit {u.unit_number}: {u.title}" for u in ai.units], "Low"
        return None, "Missing"

    else:
        return None, "Missing"


# ── Public API ────────────────────────────────────────────────────────────────

def extract_pdf(content: bytes) -> SyllabusExtraction:
    """
    Extract syllabus data from PDF bytes.

    Raises:
        EncryptedPDFError: if the PDF is password-protected.
        ScannedPDFError: if the PDF is a scanned image with no selectable text.
        ValueError: if the PDF cannot be opened at all.
    """
    try:
        import fitz
    except ImportError:
        raise ValueError("PyMuPDF (fitz) is not installed. Run: pip install pymupdf")

    try:
        doc = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise ValueError(f"Could not open PDF: {exc}") from exc

    if doc.needs_pass:
        raise EncryptedPDFError(
            "PDF is password-protected. Remove the password and re-upload."
        )

    if _is_scanned(doc):
        # Vision OCR instead of rejection; only when OCR is unavailable does
        # the scanned error surface.
        ocr_text = ocr_scanned_pdf(content)
        if not ocr_text.strip():
            raise ScannedPDFError(
                "This appears to be a scanned PDF and OCR is unavailable. "
                "Please use a digital PDF or paste the text manually."
            )
        ai = _ai_extract(ocr_text)
        if ai:
            return _ai_to_syllabus(ai)
        logger.info("AI extraction failed on OCR text — falling back to regex")
        return _regex_extract(ocr_text)

    flat_text, layout_units = _pdf_to_layout(content)

    if not flat_text.strip():
        logger.warning("Layout extraction returned empty text; falling back to page.get_text()")
        flat_text = "\n".join(page.get_text() for page in doc)

    ai = _ai_extract(flat_text, layout_units)
    if ai:
        return _ai_to_syllabus(ai)

    logger.info("AI extraction failed or skipped — falling back to regex")
    return _regex_extract(flat_text)


def extract_docx(content: bytes) -> SyllabusExtraction:
    """Extract syllabus data from DOCX bytes."""
    flat_text, layout_units = _docx_to_layout(content)

    if not flat_text.strip():
        # Final safety net
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            flat_text = "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            logger.warning("python-docx fallback failed: %s", e)

    ai = _ai_extract(flat_text, layout_units)
    if ai:
        return _ai_to_syllabus(ai)

    logger.info("AI extraction failed or skipped — falling back to regex")
    return _regex_extract(flat_text)
